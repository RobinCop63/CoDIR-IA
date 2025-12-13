import os, io, sys
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
if BASE_DIR not in sys.path:
    sys.path.append(BASE_DIR)

import streamlit as st
from datetime import datetime
from dotenv import load_dotenv
from docx import Document

from codir_engine import run_codir_session
from libre_engine import run_free_mode

load_dotenv(override=True)
st.set_page_config(page_title="Orchestrateur multi-IA v15", layout="wide")

st.title("Orchestrateur multi-IA v15")
st.caption("Mode Libre / Mode CODIR IA — OpenAI · Gemini · Claude · Mistral")

mode = st.sidebar.radio("Choisir un mode", ["Mode CODIR IA", "Mode Libre"])

with st.sidebar:
    st.header("Paramètres communs")
    temperature = st.slider("Température", 0.0, 1.0, float(os.getenv("TEMPERATURE", 0.4)))
    max_tokens = st.number_input("Max tokens", min_value=256, max_value=8000, value=int(os.getenv("MAX_TOKENS", 2000)))
    st.markdown("---")
    st.write("Modèles lus depuis `.env` :")
    st.code("\n".join([
        f"OPENAI_MODEL={os.getenv('OPENAI_MODEL','gpt-4o')}",
        f"GEMINI_MODEL={os.getenv('GEMINI_MODEL','gemini-2.5-flash')}",
        f"ANTHROPIC_MODEL={os.getenv('CLAUDE_MODEL','claude-sonnet-4-5')}",
        f"MISTRAL_MODEL={os.getenv('MISTRAL_MODEL','mistral-large-latest')}",
    ]))

if mode == "Mode CODIR IA":
    st.subheader("CODIR IA — Session hebdomadaire")
    brief = st.text_area("Brief du dirigeant (semaine précédente)", height=220,
                         value="Contexte business, actualités internes/externes, objectifs, risques/opportunités, points pour décision.")
    run_btn = st.button("Lancer la session CODIR IA", type="primary")
    if run_btn:
        with st.spinner("Génération en cours..."):
            result = run_codir_session(brief, temperature=temperature, max_tokens=max_tokens)
        st.success("Session terminée.")
        st.markdown("### Stratégie"); st.write(result["outputs"]["strategie"])
        st.markdown("### Marketing & Communication"); st.write(result["outputs"]["marketing"])
        st.markdown("### Finance & Fiscalité (FR)"); st.write(result["outputs"]["finance"])
        st.markdown("## Synthèse — Direction Générale"); st.write(result["outputs"]["direction_generale"])

        def export_docx(r: dict) -> bytes:
            doc = Document()
            doc.add_heading("CODIR IA – Compte-rendu", level=0)
            doc.add_paragraph(f"Date : {datetime.now().strftime('%d/%m/%Y')}")
            doc.add_paragraph("\nRésumé de la semaine précédente (brief):\n" + r["inputs"]["brief"])
            doc.add_heading("Stratégie", level=1); doc.add_paragraph(r["outputs"]["strategie"])
            doc.add_heading("Marketing & Communication", level=1); doc.add_paragraph(r["outputs"]["marketing"])
            doc.add_heading("Finance & Fiscalité (FR)", level=1); doc.add_paragraph(r["outputs"]["finance"])
            doc.add_heading("Direction Générale – Synthèse", level=1); doc.add_paragraph(r["outputs"]["direction_generale"])
            bio = io.BytesIO(); doc.save(bio); bio.seek(0); return bio.read()

        st.download_button("📄 Télécharger le compte-rendu (.docx)",
                           data=export_docx(result),
                           file_name=f"CODIR_IA_CR_{datetime.now().strftime('%Y-%m-%d')}.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

else:
    st.subheader("Mode Libre — tester et comparer les IA")
    system = st.text_area("Contexte / System prompt (optionnel)", height=120)
    user_prompt = st.text_area("Message", height=200, value="Bonjour à tous ! Faites une courte présentation.")
    cols = st.columns(4)
    with cols[0]: chk_openai = st.checkbox("OpenAI", True)
    with cols[1]: chk_gemini = st.checkbox("Gemini", True)
    with cols[2]: chk_claude = st.checkbox("Claude", True)
    with cols[3]: chk_mistral = st.checkbox("Mistral", True)
    run_btn = st.button("Interroger les IA sélectionnées", type="primary")
    if run_btn:
        providers = [p for p,b in [("OpenAI",chk_openai),("Gemini",chk_gemini),("Claude",chk_claude),("Mistral",chk_mistral)] if b]
        with st.spinner("Génération en cours..."):
            res = run_free_mode(user_prompt, system=system, providers=providers,
                                temperature=temperature, max_tokens=max_tokens)
        for p,out in res.items():
            st.markdown(f"### {p}")
            st.write(out)

        md = ["# Comparatif des réponses – Mode Libre", f"_Date: {datetime.now().isoformat(timespec='seconds')}_", ""]
        for p,out in res.items():
            md.append(f"## {p}\n{out}\n")
        st.download_button("📝 Télécharger les réponses (.md)",
                           data="\n".join(md).encode("utf-8"),
                           file_name=f"Libre_Comparatif_{datetime.now().strftime('%Y-%m-%d')}.md",
                           mime="text/markdown")
